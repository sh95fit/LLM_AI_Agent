# backend/services/file_parser.py
import os
import logging
import tempfile
from pathlib import Path

logger = logging.getLogger(__name__)

class FileParser:
    """PDF, Word, 이미지 파일을 텍스트로 변환하는 파서."""

    async def parse(self, file_path: str, content_type: str = "") -> str:
        """파일 경로와 MIME 타입을 받아 텍스트를 반환합니다."""
        try:
            ct = (content_type or "").lower()
            ext = Path(file_path).suffix.lower()

            if ct == "application/pdf" or ext == ".pdf":
                return self._parse_pdf(file_path)
            elif ct in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        "application/msword") or ext in (".docx", ".doc"):
                return self._parse_docx(file_path)
            elif ct.startswith("image/") or ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
                return self._parse_image(file_path)
            elif ct == "text/plain" or ext == ".txt":
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()
            else:
                # 확장자 기반 재시도
                if ext == ".pdf":
                    return self._parse_pdf(file_path)
                elif ext in (".docx", ".doc"):
                    return self._parse_docx(file_path)
                else:
                    raise ValueError(f"지원하지 않는 파일 형식: {content_type} / {ext}")
        except Exception as e:
            logger.error(f"파일 파싱 실패 ({content_type}): {e}")
            raise

    def _parse_pdf(self, file_path: str) -> str:
        """PyMuPDF로 PDF 텍스트 추출."""
        import fitz  # PyMuPDF
        result = []
        page_count = 0
        try:
            with fitz.open(file_path) as doc:
                page_count = len(doc)          # ← with 블록 안에서 저장
                for page in doc:
                    text = page.get_text("text")
                    if text.strip():
                        result.append(text.strip())
            # with 블록 밖에서는 page_count 변수만 사용
            logger.info(f"PDF 파싱 완료: {page_count}페이지, {len(result)}개 섹션")
            return "\n\n".join(result)
        except Exception as e:
            logger.error(f"PDF 파싱 에러: {e}")
            raise

    def _parse_docx(self, file_path: str) -> str:
        """python-docx로 Word 파일 텍스트 추출."""
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            # 표 내용도 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            result = "\n\n".join(paragraphs)
            logger.info(f"DOCX 파싱 완료: {len(paragraphs)}개 단락")
            return result
        except Exception as e:
            logger.error(f"DOCX 파싱 에러: {e}")
            raise

    def _parse_image(self, file_path: str) -> str:
        """pytesseract로 이미지 OCR 처리."""
        try:
            import pytesseract
            from PIL import Image
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img, lang="kor+eng")
            result = text.strip()
            logger.info(f"이미지 OCR 완료: {len(result)}자")
            return result
        except ImportError:
            logger.warning("pytesseract 미설치 – 이미지 OCR 생략")
            return ""
        except Exception as e:
            logger.error(f"이미지 파싱 에러: {e}")
            raise
